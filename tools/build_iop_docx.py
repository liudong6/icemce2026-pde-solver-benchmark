from __future__ import annotations

import csv
import math
from collections import Counter, defaultdict
from pathlib import Path
from re import split
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "template" / "icemce2026" / "WordGuidelines" / "WordGuidelines" / "IOP-ConfSer-template.docx"
OUT = ROOT / "paper" / "icemce2026_iop_manuscript.docx"
FIG_DIR = ROOT / "paper" / "figures"
RAW_DIR = ROOT / "results" / "raw"
LINE_SPACING = 1.12
EQUATION_LINE_SPACING = 1.14
EQUATION_SPACE_BEFORE_PT = 3
EQUATION_SPACE_AFTER_PT = 3


REFERENCES = [
    "Virtanen P, Gommers R, Oliphant T E, Haberland M, Reddy T, Cournapeau D et al 2020 SciPy 1.0: fundamental algorithms for scientific computing in Python Nature Methods 17 261-272. doi:10.1038/s41592-019-0686-2",
    "Harris C R, Millman K J, van der Walt S J, Gommers R, Virtanen P, Cournapeau D et al 2020 Array programming with NumPy Nature 585 357-362. doi:10.1038/s41586-020-2649-2",
    "Lam S K, Pitrou A and Seibert S 2015 Numba: a LLVM-based Python JIT compiler Proc. Second Workshop on the LLVM Compiler Infrastructure in HPC 1-6. doi:10.1145/2833157.2833162",
    "Bell N, Olson L N, Schroder J and Southworth B 2023 PyAMG: algebraic multigrid solvers in Python Journal of Open Source Software 8 5495. doi:10.21105/joss.05495",
    "Kuhn M B, Henry de Frahan M T, Mohan P, Deskos G, Churchfield M, Cheung L et al 2025 AMR-Wind: a performance-portable, high-fidelity flow solver for wind farm simulations Wind Energy 28. doi:10.1002/we.70010",
    "Li L, Fu X, Zheng X, Li H and Li J 2026 GPU-accelerated finite-element method for the three-dimensional unstructured mesh atmospheric dynamic framework Geoscientific Model Development 19 7525-7544. doi:10.5194/gmd-19-7525-2026",
    "Adams M F, Chen J and Sturdevant B 2026 Fast solvers for tokamak fluid models with PETSc Computer Physics Communications 327 110293. doi:10.1016/j.cpc.2026.110293",
    "Kumar S, Romero J, Seo J-H, Fatica M and Mittal R 2026 A GPU-accelerated sharp interface immersed boundary solver for large scale flow simulations AIAA SCITECH 2026 Forum. doi:10.2514/6.2026-0705",
    "Lapillonne X, Hupp D, Gessler F, Walser A, Pauling A, Lauber A et al 2026 Operational numerical weather prediction with ICON on GPUs (version 2024.10) Geoscientific Model Development 19 755-772. doi:10.5194/gmd-19-755-2026",
    "Strikwerda J C 2004 Finite Difference Schemes and Partial Differential Equations 2nd edn (Philadelphia: SIAM). doi:10.1137/1.9780898717938",
    "Hestenes M R and Stiefel E 1952 Methods of conjugate gradients for solving linear systems Journal of Research of the National Bureau of Standards 49 409-436. doi:10.6028/jres.049.044",
    "Saad Y 2003 Iterative Methods for Sparse Linear Systems 2nd edn (Philadelphia: SIAM). doi:10.1137/1.9780898718003",
    "Briggs W L, Henson V E and McCormick S F 2000 A Multigrid Tutorial 2nd edn (Philadelphia: SIAM). doi:10.1137/1.9780898719505",
    "Ruge J W and Stuben K 1987 Algebraic multigrid in McCormick S F (ed) Multigrid Methods (Philadelphia: SIAM) 73-130. doi:10.1137/1.9781611971057.ch4",
    "Bernaschi M, Celestini A, Richelli G and D'Ambra P 2026 On the energy efficiency of sparse matrix computations on multi-GPU clusters Future Generation Computer Systems 183 108519. doi:10.1016/j.future.2026.108519",
    "Welter A and Nguyen N C 2026 Preconditioning techniques for hybridizable discontinuous Galerkin discretizations on GPU architectures Computer Methods in Applied Mechanics and Engineering 456 118951. doi:10.1016/j.cma.2026.118951",
    "Yuan F, Yang X, Huang Y, Dong D, Xu C, Liu J et al 2025 CRAMG: a communication-reduced algebraic multigrid method Proc. 39th ACM International Conference on Supercomputing 397-411. doi:10.1145/3721145.3725764",
    "Green D, Hu X, Lore J, Mu L and Stowell M L 2022 An efficient high-order numerical solver for diffusion equations with strong anisotropy Computer Physics Communications 276 108333. doi:10.1016/j.cpc.2022.108333",
    "NVIDIA Corporation 2026 CUDA C++ Programming Guide. https://docs.nvidia.com/cuda/cuda-c-programming-guide/ accessed 3 September 2026",
    "NVIDIA Corporation 2026 Numba-CUDA Documentation. https://nvidia.github.io/numba-cuda/ accessed 3 September 2026",
    "Pekkila J, Lappi O, Robertsen F and Korpi-Lagg M J 2025 Stencil computations on AMD and Nvidia graphics processors: performance and tuning strategies Concurrency and Computation: Practice and Experience 37. doi:10.1002/cpe.70129",
    "Makhmut Y, Imankulov T, Gorlatch S and Matkerim B 2026 A CUDA performance study of global- and shared-memory kernels for the Buckley-Leverett polymer-flooding problem Applied Sciences 16 5449. doi:10.3390/app16115449",
    "Koskela T, Christidi I, Giordano M, Dubrovska E, Quinn J, Maynard C et al 2023 Principles for automated and reproducible benchmarking Proc. SC '23 Workshops 609-618. doi:10.1145/3624062.3624133",
]


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def estimate_order(h_values: list[float], errors: list[float]) -> float:
    log_h = [math.log(value) for value in h_values]
    log_e = [math.log(value) for value in errors]
    mean_h = sum(log_h) / len(log_h)
    mean_e = sum(log_e) / len(log_e)
    numerator = sum((h - mean_h) * (e - mean_e) for h, e in zip(log_h, log_e))
    denominator = sum((h - mean_h) ** 2 for h in log_h)
    return numerator / denominator


def sci_plain(value: float, precision: int = 2) -> str:
    return f"{value:.{precision}e}".replace("e-0", "e-").replace("e+0", "e+")


def method_label(method: str, threads: str | None = None) -> str:
    labels = {
        "cg": "CG",
        "jacobi-pcg": "Jacobi-PCG",
        "amg-pcg": "AMG-PCG",
        "numpy-vectorized": "NumPy",
        "numba-serial": "Numba serial",
        "cuda-kernel": "CUDA kernel",
        "numba-parallel-cpu": "CPU Numba",
        "linear_cpu_vs_cuda": "Illustrative fitted crossover",
    }
    if method == "numba-parallel":
        return f"Numba parallel, {threads} threads"
    return labels.get(method, method)


def bool_value(value: str) -> bool:
    return str(value).strip().lower() in {"1", "true", "yes"}


def convergence_table_rows() -> list[list[str]]:
    rows = read_csv_rows(RAW_DIR / "convergence.csv")
    by_case: dict[str, list[dict[str, str]]] = {}
    for row in rows:
        by_case.setdefault(row["case_name"], []).append(row)
    output: list[list[str]] = []
    for case_name, label in [
        ("smooth_constant", "Constant coefficient"),
        ("smooth_variable", "Smooth variable coefficient"),
    ]:
        case_rows = sorted(by_case[case_name], key=lambda item: int(item["n"]))
        hs = [float(row["h"]) for row in case_rows]
        l2_errors = [float(row["l2_error"]) for row in case_rows]
        linf_errors = [float(row["linf_error"]) for row in case_rows]
        finest = case_rows[-1]
        output.append(
            [
                label,
                f"{estimate_order(hs, l2_errors):.2f}",
                f"{estimate_order(hs, linf_errors):.2f}",
                sci_plain(float(finest["l2_error"])),
                sci_plain(float(finest["residual_norm"])),
            ]
        )
    return output


def interface_verification_rows() -> list[list[str]]:
    rows = read_csv_rows(RAW_DIR / "interface_verification.csv")
    grouped = {
        (int(row["n"]), row["face_average"]): row
        for row in rows
    }
    output: list[list[str]] = []
    for n in sorted({key[0] for key in grouped}):
        arithmetic = grouped[(n, "arithmetic")]
        harmonic = grouped[(n, "harmonic")]
        output.append(
            [
                str(n),
                sci_plain(float(arithmetic["l2_error"])),
                sci_plain(float(harmonic["l2_error"])),
                sci_plain(float(arithmetic["flux_error_abs"])),
                sci_plain(float(harmonic["flux_error_abs"])),
            ]
        )
    return output


def solver_table_rows() -> list[list[str]]:
    rows = read_csv_rows(RAW_DIR / "solver_benchmark.csv")
    rows_256 = [row for row in rows if int(row["n"]) == 256]
    output: list[list[str]] = []
    for case_name, label in [
        ("smooth_variable_solver", "Smooth variable"),
        ("high_contrast_solver", "High contrast"),
    ]:
        for method in ["cg", "jacobi-pcg", "amg-pcg"]:
            row = next(
                item
                for item in rows_256
                if item["case"] == case_name and item["method"] == method
            )
            output.append(
                [
                    label,
                    method_label(method),
                    row["iterations"],
                    f"{float(row['total_seconds']):.3f}",
                    sci_plain(float(row["residual_norm"])),
                ]
            )
    return output


def performance_table_rows() -> list[list[str]]:
    cpu_rows = read_csv_rows(RAW_DIR / "cpu_scaling.csv")
    gpu_rows = read_csv_rows(RAW_DIR / "gpu_stencil.csv")
    cpu_by_key = {
        (row["method"], row["threads"]): row
        for row in cpu_rows
        if int(row["n"]) == 2048
    }
    output: list[list[str]] = []
    for method, threads in [
        ("numpy-vectorized", "1"),
        ("numba-serial", "1"),
        ("numba-parallel", "4"),
    ]:
        row = cpu_by_key[(method, threads)]
        output.append(
            [
                "CPU stencil",
                method_label(method, threads),
                row["n"],
                f"{float(row['seconds_per_apply']):.4g} s/apply",
                f"{float(row['estimated_gbytes_per_second']):.2f} GB/s",
            ]
        )
    for row in sorted(
        [
            item
            for item in gpu_rows
            if item["method"] == "cuda-kernel" and int(item["n"]) in {2048, 4096}
        ],
        key=lambda item: int(item["n"]),
    ):
        output.append(
            [
                "GPU crossover",
                method_label(row["method"]),
                row["n"],
                f"{sci_plain(float(row['seconds_per_step']))} s/step",
                f"{float(row['speedup_vs_cpu']):.2f}x vs CPU",
            ]
        )
    return output


def coefficient_difficulty_rows() -> list[list[str]]:
    decision_rows = read_csv_rows(RAW_DIR / "solver_decision_map.csv")
    conditioning_rows = read_csv_rows(RAW_DIR / "conditioning.csv")
    max_conditioning_n = max(int(row["n"]) for row in conditioning_rows)
    condition_by_case = {
        row["coefficient_case"]: row
        for row in conditioning_rows
        if int(row["n"]) == max_conditioning_n
    }
    metrics_by_case: dict[str, dict[str, str]] = {}
    for row in decision_rows:
        if int(row["n"]) == 64:
            metrics_by_case[row["coefficient_case"]] = row

    grouped: dict[str, list[dict[str, str]]] = defaultdict(list)
    for row in metrics_by_case.values():
        grouped[row["family"]].append(row)

    rows: list[list[str]] = []
    for family in ["constant", "smooth", "inclusion", "layered", "checkerboard"]:
        candidates = grouped.get(family, [])
        if not candidates:
            continue
        row = max(candidates, key=lambda item: float(item["contrast_target"]))
        condition = condition_by_case.get(row["coefficient_case"])
        rows.append(
            [
                row["coefficient_case"],
                row["family"].title(),
                f"{float(row['contrast_observed']):.1f}",
                f"{float(row['grad_logk_inf']):.2f}",
                f"{float(row['total_variation_proxy']):.2f}",
                sci_plain(float(condition["condition_estimate"])) if condition is not None else "-",
            ]
        )
    return rows


def difficulty_relationship_rows() -> list[list[str]]:
    raw_rows = read_csv_rows(RAW_DIR / "difficulty_relationships.csv")
    pooled_rows = [row for row in raw_rows if row["n_group"] == "all"]
    by_key = {(row["descriptor"], row["response"]): row for row in pooled_rows}
    descriptors = [
        ("contrast_observed", "Cₖ contrast"),
        ("grad_logk_inf", "Gₖ(h) log-gradient"),
        ("total_variation_proxy", "Vₖ TV proxy"),
        ("condition_estimate", "κ(A) proxy"),
    ]

    def fixed_grid_text(descriptor: str, response: str) -> str:
        values = [
            float(row["spearman_rho"])
            for row in raw_rows
            if row["descriptor"] == descriptor
            and row["response"] == response
            and row["n_group"] != "all"
        ]
        if not values:
            return "-"
        values = sorted(values)
        midpoint = len(values) // 2
        if len(values) % 2 == 0:
            median_value = 0.5 * (values[midpoint - 1] + values[midpoint])
        else:
            median_value = values[midpoint]
        min_value = values[0]
        max_value = values[-1]
        if max_value - min_value < 0.005:
            return f"{median_value:.2f}"
        return f"{median_value:.2f} ({min_value:.2f}-{max_value:.2f})"

    descriptor_labels = {
        "contrast_observed": "Ck contrast",
        "grad_logk_inf": "Gk(h) log-gradient",
        "total_variation_proxy": "Vk TV proxy",
        "condition_estimate": "kappa(A) proxy",
    }
    rows: list[list[str]] = []
    for descriptor, label in descriptors:
        label = descriptor_labels.get(descriptor, label)
        cg = by_key.get((descriptor, "cg_iterations"))
        speedup = by_key.get((descriptor, "selected_speedup_vs_cg"))
        if cg is None and speedup is None:
            continue
        rows.append(
            [
                label,
                f"{float(cg['spearman_rho']):.2f}" if cg is not None else "-",
                fixed_grid_text(descriptor, "cg_iterations"),
                f"{float(speedup['spearman_rho']):.2f}" if speedup is not None else "-",
                fixed_grid_text(descriptor, "selected_speedup_vs_cg"),
            ]
        )
    return rows


def decision_summary_rows() -> list[list[str]]:
    rows = [row for row in read_csv_rows(RAW_DIR / "solver_decision_map.csv") if bool_value(row["is_best"])]
    grouped: dict[int, list[dict[str, str]]] = defaultdict(list)
    for row in rows:
        grouped[int(row["n"])].append(row)

    output: list[list[str]] = []
    for n in sorted(grouped):
        group = grouped[n]
        counts = Counter(row["best_method"] for row in group)
        speedups = sorted(float(row["speedup_vs_cg"]) for row in group)
        median = speedups[len(speedups) // 2]
        if len(speedups) % 2 == 0:
            median = 0.5 * (speedups[len(speedups) // 2 - 1] + speedups[len(speedups) // 2])
        output.append(
            [
                str(n),
                str(len(group)),
                str(counts.get("cg", 0)),
                str(counts.get("jacobi-pcg", 0)),
                str(counts.get("amg-pcg", 0)),
                f"{median:.2f}",
                f"{max(speedups):.2f}",
            ]
        )
    return output


def averaging_sensitivity_rows() -> list[list[str]]:
    rows = read_csv_rows(RAW_DIR / "averaging_sensitivity.csv")
    by_key = {
        (row["coefficient_case"], int(row["n"]), row["face_average"], row["method"]): row
        for row in rows
    }

    def compact_method(method: str) -> str:
        return {
            "cg": "CG",
            "jacobi-pcg": "Jacobi",
            "amg-pcg": "AMG",
        }.get(method, method_label(method))

    output: list[list[str]] = []
    for case in ["inclusion_c100", "layered_c100", "checkerboard_c100"]:
        for n in [64, 128, 256]:
            cg_a = by_key[(case, n, "arithmetic", "cg")]
            cg_h = by_key[(case, n, "harmonic", "cg")]
            jac_a = by_key[(case, n, "arithmetic", "jacobi-pcg")]
            jac_h = by_key[(case, n, "harmonic", "jacobi-pcg")]
            amg_a = by_key[(case, n, "arithmetic", "amg-pcg")]
            amg_h = by_key[(case, n, "harmonic", "amg-pcg")]
            best_a = next(
                row
                for row in rows
                if row["coefficient_case"] == case
                and int(row["n"]) == n
                and row["face_average"] == "arithmetic"
                and bool_value(row["is_best"])
            )
            best_h = next(
                row
                for row in rows
                if row["coefficient_case"] == case
                and int(row["n"]) == n
                and row["face_average"] == "harmonic"
                and bool_value(row["is_best"])
            )
            output.append(
                    [
                        case,
                        str(n),
                        f"{float(cg_a['iterations']):.0f}/{float(cg_h['iterations']):.0f}",
                        f"{float(jac_a['iterations']):.0f}/{float(jac_h['iterations']):.0f}",
                        f"{float(amg_a['iterations']):.0f}/{float(amg_h['iterations']):.0f}",
                        f"{compact_method(best_a['method'])}/{compact_method(best_h['method'])}",
                        f"{float(best_a['speedup_vs_cg']):.2f}/{float(best_h['speedup_vs_cg']):.2f}",
                    ]
            )
    return output


def timing_stability_rows() -> list[list[str]]:
    rows = read_csv_rows(RAW_DIR / "timing_stability.csv")
    preferred = [
        ("constant", 64),
        ("smooth_c30", 256),
        ("inclusion_c100", 256),
        ("layered_c100", 256),
        ("checkerboard_c100", 256),
    ]
    selected: list[dict[str, str]] = []
    for case, n in preferred:
        row = next(
            (
                item
                for item in rows
                if item["coefficient_case"] == case and int(float(item["n"])) == n
            ),
            None,
        )
        if row is not None:
            selected.append(row)
    if not selected:
        selected = sorted(rows, key=lambda item: (int(float(item["n"])), item["coefficient_case"]))

    output: list[list[str]] = []
    for row in selected:
        status = "stable" if bool_value(row["decision_stable"]) else "variable"
        output.append(
            [
                row["coefficient_case"],
                str(int(float(row["n"]))),
                method_label(row["best_method_by_median"]),
                method_label(row["best_method_vote"]),
                f"{float(row['vote_fraction']):.2f}",
                f"{float(row['median_speedup_vs_cg']):.2f}",
                f"{100.0 * float(row['selected_rel_iqr']):.1f}%",
                status,
            ]
        )
    return output


def hardware_model_rows() -> list[list[str]]:
    rows = read_csv_rows(RAW_DIR / "hardware_crossover_model.csv")
    by_component = {row["model_component"]: row for row in rows}
    cpu = by_component["cpu"]
    cuda = by_component["cuda"]
    crossover = by_component["crossover"]
    return [
        [
            "CPU Numba",
            str(int(float(cpu["observations"]))),
            sci_plain(float(cpu["beta_seconds_per_unknown"]), precision=3),
            f"{float(cpu['r2']):.3f}",
            "-",
        ],
        [
            "CUDA kernel",
            str(int(float(cuda["observations"]))),
            sci_plain(float(cuda["beta_seconds_per_unknown"]), precision=3),
            f"{float(cuda['r2']):.3f}",
            "-",
        ],
        [
            "Illustrative fitted crossover",
            "-",
            "-",
            f"{float(crossover['r2']):.3f}",
            f"{float(crossover['crossover_grid_n']):.0f}",
        ],
    ]


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def style_name(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    names = {style.name for style in doc.styles}
    return preferred if preferred in names else fallback


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(4.0)
    section.bottom_margin = Cm(2.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)

    for style in doc.styles:
        if hasattr(style, "font"):
            style.font.name = "Times New Roman"
            style.font.size = style.font.size or Pt(10)
            if style._element.rPr is not None:
                style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if hasattr(style, "paragraph_format"):
            fmt = style.paragraph_format
            fmt.line_spacing = LINE_SPACING
            if fmt.space_after is None:
                fmt.space_after = Pt(0)


def normalize_runs(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_runs_with_citations(paragraph, text: str) -> None:
    parts = split(r"(\[\[[^\]]+\]\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("[[") and part.endswith("]]"):
            citation = part[2:-2]
            run = paragraph.add_run(f"[{citation}]")
            run.font.superscript = True
            run.font.name = "Times New Roman"
        else:
            run = paragraph.add_run(part)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_paragraph(doc: Document, text: str, style: str, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    add_runs_with_citations(p, text)
    normalize_runs(p)
    return p


def add_section(doc: Document, number: int, title: str) -> None:
    p = doc.add_paragraph(style=style_name(doc, "IOP-CS-SectionHead"))
    p.add_run(f"{number}. {title}")
    normalize_runs(p)


def add_subsection(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph(style=style_name(doc, "IOP-CS-SubsectionHeading"))
    p.add_run(f"{number}. {title}")
    normalize_runs(p)


def no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def math_text(text: str) -> str:
    return f'<m:r><m:t xml:space="preserve">{escape(text)}</m:t></m:r>'


def math_expr(value: str) -> str:
    return value if value.startswith("<m:") else math_text(value)


def math_seq(*parts: str) -> str:
    return "".join(math_expr(part) for part in parts)


def math_sub(base: str, subscript: str) -> str:
    return (
        "<m:sSub><m:e>"
        f"{math_expr(base)}"
        "</m:e><m:sub>"
        f"{math_expr(subscript)}"
        "</m:sub></m:sSub>"
    )


def math_sup(base: str, superscript: str) -> str:
    return (
        "<m:sSup><m:e>"
        f"{math_expr(base)}"
        "</m:e><m:sup>"
        f"{math_expr(superscript)}"
        "</m:sup></m:sSup>"
    )


def math_frac(numerator: str, denominator: str) -> str:
    return (
        '<m:f><m:fPr><m:type m:val="bar"/></m:fPr><m:num>'
        f"{math_expr(numerator)}"
        "</m:num><m:den>"
        f"{math_expr(denominator)}"
        "</m:den></m:f>"
    )


def equation_omml_lines(text: str, number: int) -> list[str]:
    if number == 1:
        return [
            math_seq("−∇·(k(x,y)∇u(x,y)) = f(x,y),  (x,y) ∈ ", math_sup("(0,1)", "2"))
        ]
    if number == 2:
        return [math_seq("k(x,y) = 1 + 0.5 sin(2πx) sin(2πy)")]
    if number == 3:
        return [
            math_seq(
                "k(x,y) = 1 + 99 ",
                math_sub("χ", "D"),
                "(x,y)",
            ),
            math_seq(
                "D = { (x,y) : ",
                math_sup("(x−0.5)", "2"),
                " + ",
                math_sup("(y−0.5)", "2"),
                " < ",
                math_sup("0.15", "2"),
                " }",
            )
        ]
    if number == 4:
        return [
            math_seq(
                math_sub("C", "k"),
                " = ",
                math_frac("max(k)", "min(k)"),
                ",   ",
                math_sup(math_sub("G", "k"), "(h)"),
                " = ||",
                math_sub("∇", "h"),
                " log(k)||∞",
            ),
            math_seq(
                math_sub("V", "k"),
                " = Σ|",
                math_sub("Δ", "x"),
                "k| + Σ|",
                math_sub("Δ", "y"),
                "k|",
            )
        ]
    if number == 5:
        return [
            math_seq(math_sub("(Au)", "ij"), " = ", math_frac("1", math_sup("h", "2")), " ["),
            math_seq(
                math_sub("k", "i+1/2,j"),
                "(",
                math_sub("u", "ij"),
                " − ",
                math_sub("u", "i+1,j"),
                ") + ",
                math_sub("k", "i−1/2,j"),
                "(",
                math_sub("u", "ij"),
                " − ",
                math_sub("u", "i−1,j"),
                ")",
            ),
            math_seq(
                "+ ",
                math_sub("k", "i,j+1/2"),
                "(",
                math_sub("u", "ij"),
                " − ",
                math_sub("u", "i,j+1"),
                ") + ",
                math_sub("k", "i,j−1/2"),
                "(",
                math_sub("u", "ij"),
                " − ",
                math_sub("u", "i,j−1"),
                ") ]",
            ),
        ]
    if number == 6:
        return [math_seq(math_sub("u", "exact"), "(x,y) = sin(πx) sin(πy)")]
    if number == 7:
        return [
            math_seq(
                math_sup("s", "*"),
                "(c,N) = ",
                math_sub("arg min", "s"),
                " [",
                math_sub("T", "setup"),
                "(s,c,N) + ",
                math_sub("T", "solve"),
                "(s,c,N)]",
            )
        ]
    return [math_text(text)]


def equation_omml_lines(text: str, number: int) -> list[str]:
    if number == 1:
        return [
            math_seq(
                "-div(k(x,y) grad u(x,y)) = f(x,y),  (x,y) in ",
                math_sup("(0,1)", "2"),
            )
        ]
    if number == 2:
        return [math_seq("k(x,y) = 1 + 0.5 sin(2 pi x) sin(2 pi y)")]
    if number == 3:
        return [
            math_seq("k(x,y) = 1 + 99 ", math_sub("chi", "D"), "(x,y)"),
            math_seq(
                "D = { (x,y) : ",
                math_sup("(x-0.5)", "2"),
                " + ",
                math_sup("(y-0.5)", "2"),
                " < ",
                math_sup("0.15", "2"),
                " }",
            ),
        ]
    if number == 4:
        return [
            math_seq(
                math_sub("C", "k"),
                " = ",
                math_frac("max(k)", "min(k)"),
                ",   ",
                math_sup(math_sub("G", "k"), "(h)"),
                " = ",
                math_sub(math_seq("||", math_sub("grad", "h"), " log(k)||"), "inf"),
            ),
            math_seq(
                math_sub("V", "k"),
                " = sum |",
                math_sub("Delta", "x"),
                " k| + sum |",
                math_sub("Delta", "y"),
                " k|",
            ),
        ]
    if number == 5:
        return [
            math_seq(math_sub("(Au)", "ij"), " = "),
            math_seq(
                math_sub("k", "i+1/2,j"),
                math_frac(math_seq(math_sub("u", "ij"), " - ", math_sub("u", "i+1,j")), math_sup("h", "2")),
                " + ",
                math_sub("k", "i-1/2,j"),
                math_frac(math_seq(math_sub("u", "ij"), " - ", math_sub("u", "i-1,j")), math_sup("h", "2")),
            ),
            math_seq(
                "+ ",
                math_sub("k", "i,j+1/2"),
                math_frac(math_seq(math_sub("u", "ij"), " - ", math_sub("u", "i,j+1")), math_sup("h", "2")),
                " + ",
                math_sub("k", "i,j-1/2"),
                math_frac(math_seq(math_sub("u", "ij"), " - ", math_sub("u", "i,j-1")), math_sup("h", "2")),
            ),
        ]
    if number == 6:
        return [math_seq(math_sub("u", "exact"), "(x,y) = sin(pi x) sin(pi y)")]
    if number == 7:
        return [
            math_seq(
                math_sup("s", "*"),
                "(c,N) = ",
                math_sub("arg min", "s in S"),
                " [",
                math_sub("T", "setup"),
                "(s,c,N) + ",
                math_sub("T", "solve"),
                "(s,c,N)]",
            )
        ]
    return [math_text(text)]


def add_math_to_paragraph(paragraph, text: str, number: int) -> None:
    xml = f'<m:oMath {nsdecls("m")}>{equation_omml_lines(text, number)[0]}</m:oMath>'
    paragraph._p.append(parse_xml(xml))


def add_equation(doc: Document, text: str, number: int) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.allow_autofit = False
    no_borders(table)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    table.columns[0].width = Inches(0.55)
    table.columns[1].width = Inches(5.20)
    table.columns[2].width = Inches(0.55)
    spacer_cell = table.cell(0, 0)
    eq_cell = table.cell(0, 1)
    num_cell = table.cell(0, 2)
    spacer_cell.width = Inches(0.55)
    eq_cell.width = Inches(5.20)
    num_cell.width = Inches(0.55)
    lines = equation_omml_lines(text, number)
    for idx, line in enumerate(lines):
        eq_p = eq_cell.paragraphs[0] if idx == 0 else eq_cell.add_paragraph()
        eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_p.paragraph_format.line_spacing = EQUATION_LINE_SPACING
        eq_p.paragraph_format.space_before = Pt(EQUATION_SPACE_BEFORE_PT if idx == 0 else 0)
        eq_p.paragraph_format.space_after = Pt(EQUATION_SPACE_AFTER_PT if idx == len(lines) - 1 else 0)
        eq_p.paragraph_format.keep_together = True
        eq_p.paragraph_format.keep_with_next = idx < len(lines) - 1
        xml = f'<m:oMath {nsdecls("m")}>{line}</m:oMath>'
        eq_p._p.append(parse_xml(xml))
    num_p = num_cell.paragraphs[0]
    num_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    num_p.paragraph_format.line_spacing = EQUATION_LINE_SPACING
    num_p.paragraph_format.keep_together = True
    num_p.add_run(f"({number})")
    normalize_runs(num_p)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(style=style_name(doc, "IOP-CS-CaptionText"))
    p.add_run(text)
    normalize_runs(p)
    return p


def add_data_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    caption_paragraph = add_caption(doc, caption)
    caption_paragraph.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = True
            normalize_runs(paragraph)
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
    for row in rows:
        table_row = table.add_row()
        tr_pr = table_row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        cells = table_row.cells
        for cell, value in zip(cells, row):
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = True
                normalize_runs(paragraph)
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = 1.0
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)
    spacer.add_run(" ").font.size = Pt(2)


def add_figure(doc: Document, filename: str, caption: str, width_inches: float = 5.15) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(FIG_DIR / filename), width=Inches(width_inches))
    caption_paragraph = add_caption(doc, caption)
    caption_paragraph.paragraph_format.keep_together = True


def build() -> None:
    doc = Document(TEMPLATE)
    props = doc.core_properties
    props.author = "Dong Liu"
    props.last_modified_by = "Dong Liu"
    props.title = "A Coefficient-Aware Finite-Difference Benchmark for Solver Selection and CPU/GPU Stencil Scaling in Heat-Conduction Simulation"
    props.subject = "ICEMCE 2026 reproducible heat-conduction benchmark"
    props.keywords = "heat conduction; finite difference; sparse solvers; algebraic multigrid; CPU/GPU stencil scaling; reproducible benchmark"
    clear_body(doc)
    set_document_defaults(doc)

    title_style = style_name(doc, "IOP-CS-Title")
    author_style = style_name(doc, "IOP-CS-Author")
    aff_style = style_name(doc, "IOP-CS-Affiliation")
    abstract_style = style_name(doc, "IOP-CS-Abstract", aff_style)
    body_first = style_name(doc, "IOP-CS-BodyNoIndent")
    body = style_name(doc, "IOP-CS-BodyText")
    ref_style = style_name(doc, "IOP-CS-ReferenceText")

    add_paragraph(
        doc,
        "A Coefficient-Aware Finite-Difference Benchmark for Solver Selection and CPU/GPU Stencil Scaling in Heat-Conduction Simulation",
        title_style,
        WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(doc, "Dong Liu*", author_style, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "University of Nottingham Ningbo China, Ningbo 315100, China", aff_style, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "*E-mail: ssydl3@nottingham.edu.cn", aff_style, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_paragraph(
        doc,
        "Abstract. Sparse solves for variable-coefficient heat conduction are sensitive to coefficient structure, grid resolution, preconditioner setup cost, and implementation path. This paper presents a verified and reproducible finite-difference benchmark that connects these factors under controlled forcing and bounded hardware measurements. Smooth manufactured solutions recover second-order L2 convergence, while an aligned two-material test verifies discontinuous-interface flux treatment and motivates an arithmetic-versus-harmonic face-averaging sensitivity check. Across the tested coefficient families, contrast alone does not determine conditioning, and the preferred setup-inclusive solver changes with resolution: Jacobi-PCG is selected throughout the smallest single-pass decision grid, whereas the tested PyAMG smoothed-aggregation PCG configuration dominates N = 128 and N = 256. Representative repeated timings support the difficult-case selections, while the averaging sensitivity preserves the solver-class transition. CUDA results are reported only for resident-data stencil kernels and are not interpreted as full sparse-solver timings. The contribution is a controlled benchmark methodology rather than a new discretisation, Krylov method, AMG algorithm, or GPU kernel. The scripts, raw CSV files, tests, generated tables, and figures are provided as a supplementary artefact.",
        abstract_style,
    )
    doc.add_paragraph()

    add_section(doc, 1, "Introduction")
    add_paragraph(
        doc,
        "Heat transfer, diffusion, and potential-flow approximations remain central modelling components in electrical, mechanical, and computer-aided engineering. Even when the governing physics is written compactly, the numerical workload becomes a sequence of large sparse linear systems or repeated stencil updates. For short-cycle engineering studies, the scientific-computing question is therefore not only whether a discretisation converges. It is also whether the solver, preconditioner, and implementation path scale predictably under reproducible conditions.",
        body_first,
    )
    add_paragraph(
        doc,
        "The research question is: under what coefficient and resolution conditions does a stronger preconditioner become worthwhile once setup cost is included, and are those solver-selection conclusions robust to interface averaging and implementation boundaries?",
        body,
    )
    add_paragraph(
        doc,
        "This paper evaluates a compact, self-contained benchmark for the steady heat-conduction equation with spatially varying conductivity. The model is simple enough to allow manufactured-solution verification and still exposes solver behaviour under smooth, inclusion, layered, and checkerboard coefficient fields. The implementation uses Python, NumPy, SciPy, PyAMG, Numba, and Numba-CUDA, reflecting a toolchain that is accessible on a normal workstation while still supporting sparse solvers and hardware-specific kernels [[1-4]]. Recent large-scale examples motivate the same combined view: AMR-Wind and ICON-on-GPU report performance-portable engineering solvers, while GPU finite-element, tokamak-fluid, and immersed-boundary studies expose accelerator-specific PDE solver constraints [[5-9]].",
        body,
    )
    add_paragraph(
        doc,
        "Prior numerical-analysis texts establish the convergence properties of finite-difference discretisations and the role of Krylov and multigrid methods in sparse PDE solves [[10-14]]. Recent sparse and PDE solver studies further underline that preconditioning, sparse matrix execution, and accelerator mapping remain active constraints for modern workloads [[15,16]]; recent AMG and diffusion-solver work targets communication cost and coefficient difficulty more directly [[17,18]]. However, many small engineering simulation papers report only a final plot or a single runtime. That practice makes it difficult to separate discretisation error, iterative convergence, preconditioner effect, and implementation throughput. General reproducible-benchmarking work provides principles for automated reruns and evidence retention [[23]]; the narrower contribution here is to instantiate that discipline inside a PDE heat-conduction benchmark whose claims link directly to discretisation, solver selection, coefficient descriptors, and local hardware measurements. The gap considered here is deliberately bounded: one heat-conduction problem family is used to connect PDE verification, solver scaling, and CPU/GPU stencil performance.",
        body,
    )
    add_paragraph(
        doc,
        "We do not propose a new finite-difference scheme, Krylov method, AMG algorithm, or GPU kernel. The paper instead makes five scoped benchmark contributions. First, it implements a conservative finite-difference operator for -div(k grad u), verifies second-order convergence with smooth manufactured solutions, and adds an analytic jump-interface check for face-averaging behaviour. Second, it introduces a coefficient-family grid and difficulty descriptors for smooth, inclusion, layered, and checkerboard conductivity fields, with an arithmetic/harmonic sensitivity check for the discontinuous high-contrast cases. Third, it connects those descriptors to solver behaviour through pooled and fixed-grid descriptive rank correlations with fixed-grid case-resampling intervals; the contribution is the stratified diagnostic protocol, not a claim that one scalar descriptor is sufficient on its own. Fourth, it evaluates a setup-inclusive single-solve decision protocol for CG, Jacobi-preconditioned CG, and the tested PyAMG-preconditioned CG configuration, with repeated timing checks for representative cases. Fifth, it measures CPU and CUDA stencil throughput under explicit conditions and fits a local crossover interpolation. The scope is a workstation-scale structured-grid case; no industrial geometry handling, unstructured finite elements, or production-solver robustness is claimed.",
        body,
    )

    add_section(doc, 2, "Related Work")
    add_paragraph(
        doc,
        "Finite-difference methods provide a direct route from elliptic and parabolic PDEs to sparse linear systems on structured grids [[10]]. For self-adjoint diffusion operators, conservative flux forms are especially useful because face-based coefficient averaging preserves the symmetry and locality expected by Krylov methods. Once discretised, the main computational problem is sparse linear algebra rather than pointwise equation evaluation.",
        body_first,
    )
    add_paragraph(
        doc,
        "The conjugate-gradient method remains a baseline solver for symmetric positive definite systems [[11,12]]. Its performance depends strongly on spectral conditioning, which generally deteriorates under mesh refinement and can be strongly affected by coefficient contrast and geometry. Multigrid methods reduce this sensitivity by addressing error components across scales [[13]]. Algebraic multigrid extends that idea to matrix-defined problems without requiring a hand-built geometric hierarchy [[14]]. In this study, PyAMG is used as an accessible Python implementation of algebraic multigrid [[4]].",
        body,
    )
    add_paragraph(
        doc,
        "Performance portability is a separate concern. NumPy provides high-level array operations, SciPy provides sparse linear algebra, and Numba compiles numerical Python kernels to machine code [[1-3]]. CUDA exposes massively parallel GPU execution, but launch overhead and data movement can erase speedups on small problems [[19,20]]. Recent stencil and CUDA kernel studies report sensitivity to memory traffic, kernel structure, and problem size [[21,22]]. The benchmark therefore reports GPU results as kernel-only measurements for repeated updates where data stay on the device, and it keeps raw evidence for independent reruns [[23]].",
        body,
    )

    add_section(doc, 3, "Mathematical Model and Discretisation")
    add_paragraph(
        doc,
        "The model problem is the steady heat-conduction equation with homogeneous Dirichlet boundary conditions unless otherwise stated. The coefficient k(x, y) represents spatially varying thermal conductivity.",
        body_first,
    )
    add_equation(doc, "−∇·(k(x,y)∇u(x,y)) = f(x,y),  (x,y) ∈ (0,1)²", 1)
    add_paragraph(
        doc,
        "The smooth test uses a sinusoidal variable coefficient, while the high-contrast stress test uses a circular inclusion. These two cases are deliberately synthetic: they remove geometry and data-acquisition uncertainty so that the numerical mechanisms can be measured cleanly.",
        body,
    )
    add_equation(doc, "k(x,y) = 1 + 0.5 sin(2πx) sin(2πy)", 2)
    add_equation(doc, "k(x,y) = 1 + 99 χ_D(x,y),  D = { (x,y) : (x−0.5)² + (y−0.5)² < 0.15² }", 3)
    add_paragraph(
        doc,
        "For the decision-map study, this pair is expanded into a parameterised coefficient grid. The smooth family uses k = 1 + a sin(2πx) sin(2πy) with a = (Cₖ - 1)/(Cₖ + 1), giving target contrast Cₖ. The inclusion, layered, and checkerboard families use piecewise values in {1, Cₖ} with Cₖ in {10, 30, 100}, while the smooth family uses Cₖ in {3, 10, 30} and the constant case gives Cₖ = 1.",
        body,
    )
    add_paragraph(
        doc,
        "To report contrast, coefficient sharpness, and coefficient geometry separately, each coefficient field is summarised by Cₖ, the discrete maximum gradient magnitude of log(k), and a discrete total-variation proxy. For discontinuous coefficient fields, the log-gradient descriptor depends on grid spacing and is interpreted at fixed N rather than as a continuous invariant. Small-grid spectral evidence is also recorded through the estimated condition number κ(A) = λₘₐₓ(A) / λₘᵢₙ(A).",
        body,
    )
    add_equation(doc, "Cₖ = max(k)/min(k),  Gₖ(h) = ||∇h log(k)||∞,  Vₖ = Σ|Δₓk| + Σ|Δᵧk|", 4)
    add_paragraph(
        doc,
        "The domain is discretised on an N by N uniform grid with spacing h = 1/(N - 1). Interior unknowns are ordered lexicographically. For an interior node (i, j), the conservative five-point stencil is assembled by face-based arithmetic averaging of adjacent conductivity values in the main two-dimensional decision map. Arithmetic averaging is retained as the primary benchmark convention to preserve continuity with the smooth-coefficient stencil and the original coefficient-family decision grid. For discontinuous coefficient fields, harmonic face averaging is also implemented and tested as a sensitivity check because it matches the series-resistance interpretation of one-dimensional layered conduction. Both choices produce sparse symmetric positive definite systems for positive k and homogeneous Dirichlet boundaries.",
        body,
    )
    add_equation(
        doc,
        "(Au)ᵢⱼ = h⁻²[kᵢ₊₁/₂,ⱼ(uᵢⱼ-uᵢ₊₁,ⱼ)+... ]",
        5,
    )
    add_paragraph(
        doc,
        "For smooth convergence verification, the exact solution is chosen in closed form and f is generated analytically for the constant and smooth coefficient cases. Errors are reported in discrete L2 and maximum norms over interior nodes. The observed convergence rate is obtained by a least-squares fit of log(error) against log(h). A separate one-dimensional verification problem uses k(x) = k1 for x < 1/2 and k(x) = k2 for x > 1/2 with u(0) = 0 and u(1) = 1. Its analytic solution is piecewise linear and satisfies continuity of both u and k ux at the interface; this test isolates whether arithmetic and harmonic face averages respect the expected interface flux.",
        body,
    )
    add_equation(doc, "u_exact(x,y) = sin(πx) sin(πy)", 6)

    add_section(doc, 4, "Solver and Implementation Methods")
    add_paragraph(
        doc,
        "The assembled sparse matrix is solved with three methods: unpreconditioned CG, Jacobi-preconditioned CG, and algebraic-multigrid preconditioned CG. The AMG case uses the tested PyAMG default smoothed-aggregation hierarchy as a V-cycle preconditioner for SciPy CG; this should be read as one concrete AMG configuration rather than a claim about all AMG parameter choices. Manufactured-solution runs use a relative residual tolerance of 1e-11 so that solver error does not dominate discretisation error; solver-comparison runs use SciPy CG with relative tolerance 1e-8, zero absolute tolerance, the default zero initial guess, and a maximum of 12000 iterations in the decision map. All coefficient-family solver comparisons use the same deterministic unit right-hand side at a given grid size, so changes in iteration count are not confounded by case-specific forcing. The reported time separates preconditioner setup from iterative solve time and also gives the total time. Failed convergence would be recorded as a structured row in the raw CSV artefact, although all reported benchmark cases converged.",
        body_first,
    )
    add_paragraph(
        doc,
        "Matrix-free performance is measured with a separate stencil apply matching Equation (5). Three CPU paths are compared: NumPy vectorisation, Numba serial compilation, and Numba parallel compilation with controlled thread counts. Throughput is reported as an estimated stencil bandwidth using a fixed byte model: 11 double-precision values per interior point for the variable-coefficient stencil, or 88 bytes per point. The Jacobi crossover uses five double-precision values per interior point, or 40 bytes per point. These estimates are operational metrics for comparing implementations of the same stencil family, not hardware peak-bandwidth or direct DRAM-transaction claims.",
        body,
    )
    add_paragraph(
        doc,
        "The CUDA experiment uses a repeated Jacobi update rather than the full variable-coefficient operator. This isolates a simple memory-bound stencil and avoids conflating kernel throughput with sparse-solver algorithmics. CPU and GPU arrays are allocated once, warm-up iterations are run, and timing covers repeated kernel execution. Host-device transfers are excluded from the speedup calculation. This boundary matches repeated time-stepping or smoother-like updates where data remain resident on the device, but it should not be read as end-to-end application speedup.",
        body,
    )
    add_paragraph(
        doc,
        "The setup-inclusive single-solve decision protocol selects the fastest converged solver by setup-plus-solve time, making one-time AMG setup cost visible rather than treating iteration count as the only outcome. It is not a repeated-right-hand-side reuse model; such a model would require an explicit solve count m. To make the coefficient descriptors operational rather than decorative, the benchmark also reports Spearman rank correlations between Ck, Gk(h), Vk, the available-grid condition estimate, CG iterations, and the speedup of the selected solver over CG. The coefficient set is a designed benchmark grid rather than a random sample from a defined population, so rank correlations are treated as descriptive association diagnostics rather than population-level inferential estimates. Pooled rows summarise all case-size entries but are not treated as independent samples because the same coefficient family appears at multiple grid sizes. Fixed-N rows provide the safer coefficient-level reading and include deterministic case-resampling intervals. Iteration-count associations are treated as the primary solver-difficulty diagnostic; associations with selected-solver speedup are secondary and exploratory because the full decision grid uses a single timing pass per cell. Permutation p-values are archived in the CSV as diagnostics rather than used to claim pairwise descriptor separation.",
        body,
    )
    add_equation(doc, "s*(c,N) = argmin_s [T_setup(s,c,N) + T_solve(s,c,N)]", 7)
    add_paragraph(
        doc,
        "Timing stability is assessed by repeating selected decision-map cases five times and then reporting the solver chosen by median total time, the per-repeat vote fraction, and the relative interquartile range of the selected method. The full decision grid remains a single-pass benchmark, while the repeat check probes whether the most important decisions are robust or near ties.",
        body,
    )
    add_paragraph(
        doc,
        "The hardware crossover model is similarly bounded: for the repeated Jacobi kernel it fits T(n) = α + βn, where n = (N - 2)², and uses the fitted curves only to interpolate the measured CPU/CUDA equal-time region.",
        body,
    )

    add_section(doc, 5, "Experiments and Results")
    add_paragraph(
        doc,
        "All experiments were run on Windows 11 with Python 3.12.9, NumPy 2.4.6, SciPy 1.18.1, Numba 0.67.0, PyAMG 5.3.0, and Numba-CUDA 0.30.4. The machine used an AMD Ryzen 9 7940HS w/ Radeon 780M Graphics CPU with 8 physical cores and 16 logical processors, approximately 16 GB of system memory reported by the operating system as 15.2 GB usable, and an NVIDIA GeForce RTX 4060 Laptop GPU with compute capability 8.9. The exact raw CSV outputs, environment metadata, and scripts are included in the supplementary artefact.",
        body_first,
    )
    add_paragraph(
        doc,
        "The decision map covers 13 coefficient cases, three grid sizes, and three Krylov/preconditioner choices. The spectral conditioning study uses N = 32, 64, and 128. The rank-correlation analysis uses all 39 case-size decisions and three fixed-grid strata of 13 cases; condition numbers are matched to N = 64 and N = 128 where available, and the N = 128 condition estimate is used as the maximum-available case-level spectral proxy for N = 256. The repeated timing check uses five repeats for one constant case and four high-difficulty cases. The face-averaging sensitivity study reruns the three discontinuous Ck = 100 stress cases with arithmetic and harmonic face averages at N = 64, 128, and 256. The hardware crossover fit used the measured Jacobi-kernel timings at N = 512, 1024, 2048, and 4096.",
        body,
    )
    add_subsection(doc, "5.1", "Verification and Coefficient Difficulty")
    add_paragraph(
        doc,
        "The manufactured-solution study confirms the expected second-order behaviour of the finite-difference scheme. Table 1 reports fitted rates using N = 32, 64, 128, and 256 with a CG residual tolerance of 1e-11. Both the constant and smooth variable-coefficient cases produced L2 rates close to two. Figure 1 gives the same result as a log-log error plot.",
        body_first,
    )
    add_data_table(
        doc,
        "Table 1. Manufactured-solution convergence for the finite-difference discretisation.",
        ["Case", "p_L2", "p_inf", "Finest L2", "Residual"],
        convergence_table_rows(),
    )
    add_figure(
        doc,
        "convergence_l2.png",
        "Figure 1. Manufactured-solution convergence of the finite-difference operator.",
    )
    add_paragraph(
        doc,
        "The discontinuous-interface check in Table 2 isolates a different error mechanism. For an aligned one-dimensional two-material conduction problem, harmonic face averaging reproduced the analytic interface flux to roundoff, while arithmetic averaging showed a grid-refined interface error. This does not make harmonic averaging a universal choice for all multidimensional discontinuities, but it prevents the benchmark from relying only on smooth manufactured-solution evidence.",
        body,
    )
    add_data_table(
        doc,
        "Table 2. Discontinuous-interface verification on a one-dimensional two-material heat-conduction problem. Harmonic face averaging recovers the analytic flux continuity to roundoff in this aligned-interface test.",
        ["N", "L2 arithmetic", "L2 harmonic", "Flux err. arithmetic", "Flux err. harmonic"],
        interface_verification_rows(),
    )
    add_paragraph(
        doc,
        "Table 3 reports representative coefficient descriptors. The spectral evidence confirms that contrast alone is not a complete difficulty measure. At N = 128, the inclusion case with Ck = 100 had an estimated condition number of 6.43e5, while the checkerboard case with the same contrast had a much smaller estimate of 3.12e4 under this discretisation. Figure 2 shows the same effect across the measured coefficient grid.",
        body,
    )
    add_data_table(
        doc,
        "Table 3. Representative coefficient-difficulty descriptors. Metrics use N = 64 for the discrete descriptors; kappa(A) is the largest available spectral estimate, here N = 128.",
        ["Case", "Family", "Ck", "Gk(h)", "TV proxy", "kappa(A)"],
        coefficient_difficulty_rows(),
    )
    add_figure(
        doc,
        "conditioning.png",
        "Figure 2. Estimated condition number response to coefficient contrast and coefficient geometry.",
    )
    add_paragraph(
        doc,
        "Table 4 turns the coefficient descriptors into a solver-behaviour diagnostic and also shows why a single pooled ranking is unsafe. Across all 39 case-size entries, the matched condition estimate had the largest association with selected-solver speedup (rho = 0.94), while Gk(h) had a larger pooled association with CG iterations (rho = 0.82) than contrast (rho = 0.63). Within fixed-N strata, the condition estimate, contrast, and Gk(h) were all positively associated with CG iterations, with median stratum correlations of rho = 0.96, 0.89, and 0.83, respectively. However, the fixed-grid samples contain only 13 coefficient cases and the case-resampling intervals overlap in the raw CSV archive. The result is therefore used as a descriptive ordering signal and as evidence for stratified reporting, not as a formal claim that these descriptors are significantly separated.",
        body,
    )
    add_data_table(
        doc,
        "Table 4. Descriptive Spearman rank correlations linking coefficient-difficulty descriptors to solver behaviour. Pooled columns use all 39 case-size entries and are not treated as independent population samples; fixed-N columns summarise the stratum-specific rank coefficients by median and, when they differ, range. Raw stratum values and case-resampling uncertainty intervals are archived in the supplementary CSV.",
        ["Descriptor", "CG pooled", "CG fixed-N", "Speedup pooled", "Speedup fixed-N"],
        difficulty_relationship_rows(),
    )

    add_subsection(doc, "5.2", "Solver Decision Map and Preconditioner Effect")
    add_paragraph(
        doc,
        "The single-pass setup-inclusive decision map gives a sharper result than a single solver comparison. Over the 13 coefficient cases, Jacobi-PCG was fastest in every N = 64 run, while AMG-PCG was fastest in every N = 128 and N = 256 run. Table 5 summarises this transition. The median speedup of the selected solver over CG increased from 2.79x at N = 64 to 7.80x at N = 256, and the largest speedup reached 25.72x for the layered Ck = 100 case. Because each decision-map cell uses one timing pass, this map is interpreted together with the repeated representative cases below.",
        body_first,
    )
    add_data_table(
        doc,
        "Table 5. Setup-inclusive single-solve decision summary over the coefficient-family grid. Counts report the fastest converged method by setup-plus-solve time from one timing pass per cell.",
        ["N", "Cases", "CG", "Jacobi-PCG", "AMG-PCG", "Median speedup", "Max speedup"],
        decision_summary_rows(),
    )
    add_figure(
        doc,
        "solver_decision_map.png",
        "Figure 3. Fastest converged solver by setup-plus-solve time over the coefficient-family grid.",
    )
    add_paragraph(
        doc,
        "Table 6 tests whether the discontinuous high-contrast solver conclusions are an artefact of arithmetic face averaging. Replacing arithmetic by harmonic averaging changed condition estimates and iteration counts, especially for checkerboard fields, but the fastest solver class was unchanged in all nine tested case-size cells: Jacobi-PCG remained fastest at N = 64, and AMG-PCG remained fastest at N = 128 and N = 256. The stronger claim is not that averaging is irrelevant, but that the reported solver-decision transition survived this targeted interface-discretisation perturbation.",
        body,
    )
    add_data_table(
        doc,
        "Table 6. Arithmetic/harmonic face-averaging sensitivity for discontinuous Ck = 100 stress cases. Each paired entry reports arithmetic/harmonic values. The full CSV also records kappa(A) using matched N for 64 and 128 and the N = 128 proxy for N = 256.",
        ["Case", "N", "CG it.", "Jacobi it.", "AMG it.", "Best solver", "Best speedup"],
        averaging_sensitivity_rows(),
    )
    add_paragraph(
        doc,
        "The repeat timing check in Table 7 gives a stricter interpretation of the map. For the high-difficulty representative cases at N = 256, AMG-PCG remained the median-best solver in every repeat, with selected-method relative IQR between 0.5% and 6.6%. The constant-coefficient N = 64 case was a near tie: Jacobi-PCG was best by median time, but the per-repeat vote favoured CG in four of five repeats and the median speedup was only 1.03x. The benchmark therefore treats that small-grid constant case as a timing boundary rather than a robust preconditioner preference.",
        body,
    )
    add_data_table(
        doc,
        "Table 7. Repeated timing stability check. Rows report the solver selected by median total time across repeats, the per-repeat vote fraction, and the relative interquartile range of the selected method.",
        ["Case", "N", "Median best", "Vote best", "Vote", "Speedup", "Rel. IQR", "Status"],
        timing_stability_rows(),
    )
    add_paragraph(
        doc,
        "Solver behaviour changed substantially when the coefficient field became high contrast. At N = 256, unpreconditioned CG required 4370 iterations in the high-contrast case, compared with 770 iterations in the smooth case. Jacobi preconditioning reduced the high-contrast iteration count to 535, while AMG-PCG reduced it to 12. The timing reduction was also large: total time fell from 6.155 s for CG to 0.288 s for AMG-PCG in the high-contrast case. Table 8 summarises the finest-grid comparison, and Figures 4 and 5 show scaling across grid sizes.",
        body,
    )
    add_data_table(
        doc,
        "Table 8. Solver comparison at N = 256 grid points in each direction.",
        ["Case", "Method", "Iter.", "Time (s)", "Rel. residual"],
        solver_table_rows(),
    )
    add_figure(
        doc,
        "solver_runtime.png",
        "Figure 4. Total setup plus solve time for CG, Jacobi-PCG, and AMG-PCG.",
    )
    add_figure(
        doc,
        "solver_iterations.png",
        "Figure 5. Iteration scaling for the solver comparison. AMG-PCG keeps iteration counts nearly flat over the tested grid sizes for these cases.",
    )

    add_subsection(doc, "5.3", "CPU and CUDA Stencil Scaling")
    add_paragraph(
        doc,
        "The matrix-free CPU stencil benchmark shows that implementation path can dominate the cost of repeated PDE stencil operations. At N = 2048, NumPy required 0.1813 s per apply, while Numba serial required 0.005325 s and Numba parallel with four threads required 0.003725 s. Increasing thread count beyond four did not improve the largest case in this run, which is consistent with bandwidth saturation and thread-management overhead rather than a numerical failure.",
        body_first,
    )
    add_paragraph(
        doc,
        "The CUDA Jacobi benchmark passed the predefined inclusion rule. At N = 512, GPU and CPU kernel times were essentially equal. At N = 1024, CUDA reached 1.52x speedup. At N = 2048 and N = 4096, the speedups increased to 3.23x and 3.51x, respectively. These results justify including GPU scaling as a bounded kernel-level result. They do not imply the same speedup for an end-to-end sparse solver.",
        body,
    )
    add_data_table(
        doc,
        "Table 9. Representative stencil-throughput measurements. CPU entries use the variable-coefficient operator at N = 2048; CUDA entries are kernel-only Jacobi steps with arrays resident on the device.",
        ["Experiment", "Method", "N", "Time", "Throughput/speedup"],
        performance_table_rows(),
    )
    add_figure(
        doc,
        "cpu_scaling.png",
        "Figure 6. Estimated CPU stencil bandwidth for the variable-coefficient stencil apply.",
    )
    add_figure(
        doc,
        "gpu_crossover.png",
        "Figure 7. CUDA kernel-only speedup for repeated Jacobi updates relative to a four-thread Numba CPU baseline.",
    )
    add_paragraph(
        doc,
        "The local linear crossover fit gives a compact way to report this boundary. Table 10 gives the fitted slopes and equal-time estimate. The fitted crossover is an illustrative interpolation near N = 755; the directly measured evidence is that N = 512 is near equal time and N = 1024 is already faster on CUDA. This fit is used as an interpolation summary over the measured range, not as a claim about hardware launch overhead outside the experiment.",
        body,
    )
    add_data_table(
        doc,
        "Table 10. Empirical linear CPU/CUDA kernel crossover model fitted over the measured Jacobi-stencil range.",
        ["Component", "Obs.", "beta (s/unknown)", "R2", "Crossover N"],
        hardware_model_rows(),
    )

    add_section(doc, 6, "Discussion")
    add_paragraph(
        doc,
        "The main numerical result is the decision-level transition rather than the superiority of a single method in isolation. At small grids, Jacobi-PCG can win because its setup cost is negligible, but the repeated timing check shows that this statement has a near-tie boundary for the constant N = 64 case. Once the grid reaches the tested N = 128 and N = 256 levels, AMG-PCG wins consistently in the decision grid and in the repeated representative checks because its solve-time reduction outweighs the one-time AMG setup cost. This behaviour is consistent with the multilevel purpose of AMG: reduce error across scales rather than relying on local diagonal rescaling.",
        body_first,
    )
    add_paragraph(
        doc,
        "The coefficient-difficulty results also show why neither contrast nor any other single descriptor should be used alone. Inclusion and checkerboard fields can share both nominal contrast and discrete log-gradient values while having very different condition estimates, because geometry changes the operator spectrum. The pooled rank correlations are useful for summarising the combined grid-size and coefficient-field effect, but the fixed-grid correlations give the safer coefficient-level reading: contrast, discrete sharpness, and spectral conditioning capture different parts of the solver response, and the small sample size prevents a strong significance claim about their pairwise ordering. This is the main methodological lesson of the descriptor analysis: benchmark descriptors should be reported with stratification and uncertainty, not collapsed into a single universal difficulty score.",
        body,
    )
    add_paragraph(
        doc,
        "The performance results add a second layer. Compiled stencil kernels are much faster than NumPy for repeated local updates, but parallel CPU scaling is not monotone. On the largest CPU stencil case, four threads outperformed eight and sixteen threads. The safest interpretation is stencil-bandwidth and scheduling saturation under this memory-bound kernel. The CUDA result is also size-dependent. It becomes useful only after the grid is large enough to offset launch overhead, and the reported speedup assumes data residency on the GPU. The fitted crossover is therefore a reporting device for this specific kernel and workstation, not a general GPU superiority claim.",
        body,
    )
    add_paragraph(
        doc,
        "The benchmark has clear boundaries. It uses a two-dimensional structured grid, synthetic coefficient fields, and workstation-scale experiments. Harmonic averaging is not a substitute for an interface-fitted or immersed-interface discretisation when curved material boundaries cut the Cartesian grid. The benchmark does not claim industrial geometry handling, unstructured finite elements, transient multiphysics coupling, or production solver robustness. Its value is instead methodological: it gives a reproducible bridge from PDE discretisation to solver choice and hardware execution, with each claim tied to a raw CSV table, a script, and, for the descriptor analysis, explicit uncertainty checks.",
        body,
    )

    add_section(doc, 7, "Reproducibility Artefact")
    add_paragraph(
        doc,
        "The supplementary artefact contains the scripts, raw CSV files, generated tables and figures, metadata, and tests needed to reproduce the reported benchmark. The main scripts cover convergence, solver comparison, the coefficient-aware decision map, conditioning, interface verification, descriptor analysis, arithmetic/harmonic sensitivity, timing stability, CPU/CUDA stencil scaling, crossover fitting, and table/figure regeneration. Unit tests cover the numerical operators, solver behaviour, summaries, plotting, tables, descriptors, conditioning, timing stability, and crossover fit; every reported manuscript number traces to a raw CSV file.",
        body_first,
    )

    add_section(doc, 8, "Conclusion")
    add_paragraph(
        doc,
        "This paper presented a coefficient-aware finite-difference benchmark for heat-conduction simulation. Smooth verification, an aligned-interface check, coefficient descriptors, setup-inclusive solver decisions, repeated timings, and bounded CPU/CUDA stencil measurements were linked to raw CSV evidence. The main lesson is methodological: solver choice should be reported with discretisation checks, coefficient stratification, setup cost, timing stability, and explicit hardware boundaries. In this benchmark, Jacobi-PCG was selected on the smallest single-pass grid, AMG-PCG dominated the tested larger grids, and CUDA speedups were meaningful only for resident-data stencil kernels at sufficiently large problem sizes.",
        body_first,
    )

    p = doc.add_paragraph(style=style_name(doc, "IOP-CS-SectionHead"))
    p.add_run("References")
    normalize_runs(p)
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph(style=ref_style)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(f"[{i}] {ref}")
        normalize_runs(p)
        for run in p.runs:
            run.font.size = Pt(8.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
